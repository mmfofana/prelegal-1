"""Service for generating DOCX documents using python-docx."""
from __future__ import annotations

import io
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from document_registry import DOCUMENT_REGISTRY

MARKDOWN_TEMPLATES_DIR = Path(__file__).parent.parent.parent / "templates"

_NAVY = RGBColor(3, 33, 71)      # #032147
_BLUE = RGBColor(32, 157, 215)   # #209dd7


def _format_date(iso_date: str) -> str:
    try:
        d = datetime.strptime(iso_date, "%Y-%m-%d")
        return f"{d.strftime('%B')} {d.day}, {d.year}"
    except ValueError:
        return iso_date


def _strip_markdown(text: str) -> str:
    """Convert basic markdown to plain text for DOCX body."""
    # Remove HTML tags from markdown-converted text
    text = re.sub(r"<[^>]+>", "", text)
    # Remove markdown bold/italic markers
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"__(.+?)__", r"\1", text)
    text = re.sub(r"_(.+?)_", r"\1", text)
    # Remove markdown links
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    return text


def _read_markdown_template(filename: str) -> str:
    path = MARKDOWN_TEMPLATES_DIR / filename
    return path.read_text(encoding="utf-8")


def _set_heading_style(paragraph, color: RGBColor, size_pt: int = 14) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = color
        run.font.size = Pt(size_pt)
        run.font.bold = True


def generate_docx(data: dict) -> bytes:
    document_type = data["document_type"]
    doc_def = DOCUMENT_REGISTRY.get(document_type)
    if doc_def is None:
        raise ValueError(f"Unknown document type: {document_type}")

    doc = Document()

    # ── Title ───────────────────────────────────────────────────────────────
    title_para = doc.add_heading(doc_def.display_name, level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_heading_style(title_para, _NAVY, 18)

    doc.add_paragraph(
        "This agreement consists of this Cover Page and the "
        f"Common Paper {doc_def.display_name} Standard Terms."
    ).runs[0].font.size = Pt(10)

    doc.add_paragraph("")

    # ── Cover Page ──────────────────────────────────────────────────────────
    cover_heading = doc.add_heading("Cover Page", level=2)
    _set_heading_style(cover_heading, _NAVY, 14)

    effective_date = _format_date(data.get("effective_date", ""))
    governing_law = data.get("governing_law", "")
    jurisdiction = data.get("jurisdiction", "")
    party1 = data.get("party1", {})
    party2 = data.get("party2", {})
    extra = data.get("extra_fields", {})

    # Extra fields
    for field_def in doc_def.extra_fields:
        key = field_def.key
        if key in ("mnda_term_years", "term_of_confidentiality_years"):
            continue
        value = extra.get(key, "")
        p = doc.add_paragraph()
        run_label = p.add_run(f"{field_def.label}: ")
        run_label.bold = True
        run_label.font.size = Pt(10)
        p.add_run(value or "[Not specified]").font.size = Pt(10)

    # Common fields
    for label, value in [
        ("Effective Date", effective_date),
        ("Governing Law", governing_law),
        (doc_def.jurisdiction_label, jurisdiction),
    ]:
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        run_label.font.size = Pt(10)
        p.add_run(value or "[Not specified]").font.size = Pt(10)

    doc.add_paragraph("")

    # Parties table
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = ""
    hdr[1].text = doc_def.party1_label.upper()
    hdr[2].text = doc_def.party2_label.upper()
    for cell in hdr:
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "032147")
        tc_pr.append(shd)

    for field_key, field_label in [
        ("company", "Company"),
        ("name", "Print Name"),
        ("title", "Title"),
        ("address", "Notice Address"),
    ]:
        row = table.add_row().cells
        row[0].text = field_label
        row[1].text = party1.get(field_key, "")
        row[2].text = party2.get(field_key, "")

    # Signature rows
    for label in ["Signature", "Date"]:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = ""
        row[2].text = ""

    doc.add_page_break()

    # ── Standard Terms (plain text from markdown) ───────────────────────────
    terms_heading = doc.add_heading("Standard Terms", level=2)
    _set_heading_style(terms_heading, _NAVY, 14)

    md_content = _read_markdown_template(doc_def.markdown_template)
    for line in md_content.splitlines():
        stripped = _strip_markdown(line).strip()
        if not stripped:
            continue
        if line.startswith("## "):
            h = doc.add_heading(stripped, level=3)
            _set_heading_style(h, _NAVY, 12)
        elif line.startswith("# "):
            h = doc.add_heading(stripped, level=2)
            _set_heading_style(h, _NAVY, 13)
        else:
            doc.add_paragraph(stripped).runs[0].font.size = Pt(9) if stripped else None

    # ── Footer ──────────────────────────────────────────────────────────────
    doc.add_paragraph("")
    footer_para = doc.add_paragraph(
        f"Common Paper {doc_def.display_name} · Free to use under CC BY 4.0"
    )
    footer_para.runs[0].font.size = Pt(8)
    footer_para.runs[0].font.color.rgb = RGBColor(136, 136, 136)

    # Serialize to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
